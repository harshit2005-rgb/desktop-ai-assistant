"""FastMCP document server for safe local document reading and search."""

from __future__ import annotations

import csv
import json
import logging
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

from fastmcp import FastMCP

logger = logging.getLogger(__name__)

mcp = FastMCP("desktop-documents")

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".json", ".py"}
TEXT_EXTENSIONS = {".txt", ".md", ".py"}
SEARCH_ROOTS = (
    Path.cwd(),
    Path.home() / "Desktop",
    Path.home() / "Documents",
    Path.home() / "Downloads",
)
MAX_READ_BYTES = 5 * 1024 * 1024
MAX_RETURN_CHARS = 120_000
MAX_CSV_ROWS = 500
SEARCH_CONTEXT_LINES = 2
MAX_SEARCH_MATCHES = 50


def _serialize_error(message: str, **extra: Any) -> dict[str, Any]:
    """Return a consistent error payload for MCP tool responses."""
    payload: dict[str, Any] = {"success": False, "error": message}
    payload.update(extra)
    return payload


def _metadata(path: Path, content: str = "", truncated: bool = False) -> dict[str, Any]:
    """Return document metadata."""
    stat = path.stat()
    return {
        "name": path.name,
        "path": str(path),
        "extension": path.suffix.lower(),
        "size_bytes": stat.st_size,
        "created": datetime.fromtimestamp(stat.st_ctime).isoformat(timespec="seconds"),
        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
        "characters_returned": len(content),
        "line_count": content.count("\n") + 1 if content else 0,
        "truncated": truncated,
    }


def _normalize_path(path: str) -> Path:
    """Resolve user paths and common bare filenames."""
    candidate = Path(path).expanduser()
    if candidate.exists() or candidate.is_absolute():
        return candidate

    query = path.strip()
    if not query:
        return candidate

    exact_matches: list[Path] = []
    fuzzy_matches: list[Path] = []
    query_lower = query.lower()
    for root in SEARCH_ROOTS:
        if not root.exists():
            continue
        try:
            for found in root.rglob("*"):
                if not found.is_file():
                    continue
                name_lower = found.name.lower()
                if name_lower == query_lower:
                    exact_matches.append(found)
                elif query_lower in name_lower:
                    fuzzy_matches.append(found)
        except OSError as exc:
            logger.debug("Unable to search %s while resolving %s: %s", root, path, exc)

    matches = exact_matches or fuzzy_matches
    if not matches:
        return candidate
    return max(matches, key=lambda item: item.stat().st_mtime)


def _trim_content(content: str) -> tuple[str, bool]:
    """Bound content returned to the agent."""
    if len(content) <= MAX_RETURN_CHARS:
        return content, False
    return content[:MAX_RETURN_CHARS], True


def _read_text(path: Path) -> str:
    """Read a text-like file with encoding fallbacks."""
    raw = path.read_bytes()[:MAX_READ_BYTES]
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _read_json(path: Path) -> str:
    """Read JSON as formatted text when possible."""
    text = _read_text(path)
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        return text
    return json.dumps(parsed, indent=2, ensure_ascii=False)


def _read_csv(path: Path) -> str:
    """Read CSV into a compact markdown-style table preview."""
    text = _read_text(path)
    sample = text.splitlines()
    reader = csv.reader(sample)
    rows = [row for _, row in zip(range(MAX_CSV_ROWS + 1), reader)]
    if not rows:
        return ""

    output: list[str] = []
    header = rows[0]
    output.append(f"CSV columns ({len(header)}): {', '.join(header)}")
    output.append(f"Rows returned: {max(0, len(rows) - 1)}")
    output.append("")
    output.append(",".join(header))
    for row in rows[1:]:
        output.append(",".join(row))

    if len(sample) > len(rows):
        output.append("")
        output.append(f"[Truncated CSV preview after {MAX_CSV_ROWS} data rows]")
    return "\n".join(output)


def _read_pdf(path: Path) -> str:
    """Extract text from a PDF using pypdf when installed."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF support requires pypdf. Install dependencies with pip install -r requirements.txt") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        if sum(len(item) for item in pages) >= MAX_RETURN_CHARS:
            break
        text = page.extract_text() or ""
        pages.append(f"\n--- Page {index} ---\n{text.strip()}")
    return "\n".join(pages).strip()


def _read_docx(path: Path) -> str:
    """Extract text from a DOCX using python-docx when installed."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "DOCX support requires python-docx. Install dependencies with pip install -r requirements.txt"
        ) from exc

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"\nTable {table_index}:")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_file_impl(path: str) -> dict[str, Any]:
    """Read supported local document content and return text plus metadata."""
    target = _normalize_path(path)
    if not target.exists():
        return _serialize_error(f"File does not exist: {target}")
    if not target.is_file():
        return _serialize_error(f"Path is not a file: {target}")

    extension = target.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        return _serialize_error(
            f"Unsupported file type: {extension or 'no extension'}",
            supported_formats=sorted(SUPPORTED_EXTENSIONS),
        )

    try:
        if target.stat().st_size > MAX_READ_BYTES and extension in TEXT_EXTENSIONS | {".json", ".csv"}:
            logger.info("Reading first %s bytes from large file %s", MAX_READ_BYTES, target)

        if extension in TEXT_EXTENSIONS:
            content = _read_text(target)
        elif extension == ".json":
            content = _read_json(target)
        elif extension == ".csv":
            content = _read_csv(target)
        elif extension == ".pdf":
            content = _read_pdf(target)
        elif extension == ".docx":
            content = _read_docx(target)
        else:
            return _serialize_error(f"Unsupported file type: {extension}")
    except Exception as exc:  # noqa: BLE001 - parsers raise library-specific exceptions.
        logger.exception("Unable to read document %s", target)
        return _serialize_error(str(exc), path=str(target))

    content, truncated = _trim_content(content)
    metadata = _metadata(target, content, truncated)
    return {"success": True, "content": content, "metadata": metadata}


def search_file_content_impl(path: str, query: str) -> dict[str, Any]:
    """Search inside a supported file and return matching lines with context."""
    if not query.strip():
        return _serialize_error("query is required")

    read_result = read_file_impl(path)
    if not read_result.get("success"):
        return read_result

    content = read_result.get("content", "")
    lines = content.splitlines()
    query_lower = query.lower()
    matches: list[dict[str, Any]] = []

    for index, line in enumerate(lines):
        if query_lower not in line.lower():
            continue

        start = max(0, index - SEARCH_CONTEXT_LINES)
        end = min(len(lines), index + SEARCH_CONTEXT_LINES + 1)
        matches.append(
            {
                "line_number": index + 1,
                "line": line,
                "context": [
                    {"line_number": context_index + 1, "text": lines[context_index]}
                    for context_index in range(start, end)
                ],
            }
        )
        if len(matches) >= MAX_SEARCH_MATCHES:
            break

    return {
        "success": True,
        "query": query,
        "path": read_result["metadata"]["path"],
        "metadata": read_result["metadata"],
        "count": len(matches),
        "matches": matches,
        "truncated_matches": len(matches) >= MAX_SEARCH_MATCHES,
    }


def summarize_file_impl(path: str, mode: str = "standard") -> dict[str, Any]:
    """Return a deterministic local summary preview for standalone MCP use."""
    read_result = read_file_impl(path)
    if not read_result.get("success"):
        return read_result

    content = read_result.get("content", "")
    summary = _extractive_summary(content, mode)
    return {
        "success": True,
        "path": read_result["metadata"]["path"],
        "metadata": read_result["metadata"],
        "mode": mode,
        "summary": summary,
    }


def _extractive_summary(content: str, mode: str) -> dict[str, Any]:
    """Create a simple local summary when no LLM is available."""
    paragraphs = [item.strip() for item in content.split("\n\n") if item.strip()]
    sentences = []
    for paragraph in paragraphs:
        for sentence in paragraph.replace("\n", " ").split(". "):
            cleaned = sentence.strip()
            if cleaned:
                sentences.append(cleaned if cleaned.endswith(".") else f"{cleaned}.")

    words = [
        word.strip(".,:;!?()[]{}\"'").lower()
        for word in content.split()
        if len(word.strip(".,:;!?()[]{}\"'")) > 4
    ]
    common = [word for word, _ in Counter(words).most_common(10)]
    selected = sentences[:5] if sentences else paragraphs[:3]
    if "one paragraph" in mode.lower():
        selected = selected[:3]

    return {
        "summary": " ".join(selected)[:2000] or "No readable text found.",
        "key_points": selected[:5],
        "important_terms": common,
        "action_items": [
            line.strip()
            for line in content.splitlines()
            if any(marker in line.lower() for marker in ("todo", "action", "next step", "follow up", "deadline"))
        ][:10],
        "conclusion": selected[-1] if selected else "No conclusion could be inferred from the readable text.",
    }


@mcp.tool()
def read_file(path: str) -> dict[str, Any]:
    """Read supported local document content and return text plus metadata."""
    return read_file_impl(path)


@mcp.tool()
def summarize_file(path: str, mode: str = "standard") -> dict[str, Any]:
    """Summarize a supported local document with a local extractive preview."""
    return summarize_file_impl(path, mode)


@mcp.tool()
def search_file_content(path: str, query: str) -> dict[str, Any]:
    """Search inside a supported file and return matching lines with context."""
    return search_file_content_impl(path, query)


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    mcp.run()
